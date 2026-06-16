#!/usr/bin/env python
"""Preprocess local model-output Word documents for Fig. 3 source tables."""

from __future__ import annotations

import argparse
import csv
import re
import sys
import traceback
from collections import Counter, defaultdict
from dataclasses import dataclass
from difflib import SequenceMatcher
from itertools import combinations
from pathlib import Path
from typing import Iterable


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

try:
    from ncfigs.constants import DIMENSION_DISPLAY, DIMENSION_FIELDS
except Exception:
    DIMENSION_FIELDS = [
        "color_change",
        "shape_change",
        "surface_texture_change",
        "volume_change",
        "dissolution_speed_time",
        "physical_state_change",
        "dissolution_medium",
        "fragment_distribution_density",
    ]
    DIMENSION_DISPLAY = {
        "color_change": "Color change",
        "shape_change": "Shape change",
        "surface_texture_change": "Surface texture change",
        "volume_change": "Area/size change",
        "dissolution_speed_time": "Dissolution speed over time",
        "physical_state_change": "Physical state change",
        "dissolution_medium": "Dissolution medium",
        "fragment_distribution_density": "Fragment distribution density",
    }


MODEL_BY_SOURCE_FOLDER = {
    "chat gpt": ("gpt-5-mini", "GPT-5-mini", "inferred_from_folder_chat_gpt"),
    "dis gpt old": ("dis-gpt", "DIS GPT", "explicit_user_mapping_dis_gpt"),
    "kimi": ("kimi-k2.5", "Kimi-K2.5", "inferred_from_folder_kimi"),
    "qwen": ("qwen3.6-plus", "Qwen3.6-plus", "inferred_from_folder_qwen"),
    "zhipu": ("glm-4.6v", "GLM-4.6V", "inferred_from_folder_zhipu"),
}

RUN_COLUMNS = [
    "model",
    "model_display_name",
    "drug",
    "run_id",
    "source_folder",
    "source_file",
    "source_id",
    "comparison_id",
    *DIMENSION_FIELDS,
    "completed_dimensions",
    "completion_rate",
    "bertscore_f1",
    "output_text",
    "reference_type",
    "notes",
]

PAIRWISE_COLUMNS = [
    "model",
    "model_display_name",
    "drug",
    "comparison_id",
    "run_id_1",
    "run_id_2",
    "source_file_1",
    "source_file_2",
    "pairwise_similarity",
    "similarity_method",
    "notes",
]

LONG_COLUMNS = [
    "model",
    "model_display_name",
    "drug",
    "run_id",
    "source_folder",
    "source_file",
    "dimension",
    "dimension_display_name",
    "present",
    "content",
    "content_length",
]

QC_COLUMNS = [
    "model",
    "model_display_name",
    "drug",
    "run_id",
    "source_folder",
    "source_file",
    "source_id",
    "parse_status",
    "completed_dimensions",
    "completion_rate",
    "missing_dimensions",
    "bertscore_f1",
    "output_text_length",
    "qc_flag",
    "warning_message",
]

MAPPING_COLUMNS = [
    "source_folder",
    "source_file",
    "model",
    "model_display_name",
    "run_id",
    "source_id",
    "include_for_fig3",
    "notes",
]


@dataclass
class SourceMeta:
    source_folder: str
    source_file: str
    source_id: str
    comparison_id: str
    run_id: str
    replicate_id: str
    model: str
    model_display_name: str
    mapping_note: str
    mapping_uncertain: bool


def ensure_parent(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    return path


def write_csv(path: Path, rows: list[dict[str, object]], columns: list[str]) -> None:
    ensure_parent(path)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=columns, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({col: row.get(col, "") for col in columns})


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def normalize_quotes(text: str) -> str:
    replacements = {
        "\u201c": '"',
        "\u201d": '"',
        "\u2018": "'",
        "\u2019": "'",
        "\u2011": "-",
        "\u2012": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\u3000": " ",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text


def read_docx(path: Path) -> tuple[str, str, str]:
    try:
        from docx import Document
    except Exception as exc:
        return "", "dependency_missing", f"python-docx is unavailable: {exc}"

    try:
        document = Document(str(path))
        parts: list[str] = []
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts), "read_ok", ""
    except Exception as exc:
        return "", "read_failed", f"{exc.__class__.__name__}: {exc}"


def infer_model(source_folder: str) -> tuple[str, str, str, bool]:
    key = source_folder.strip().lower()
    if key in MODEL_BY_SOURCE_FOLDER:
        model, display, note = MODEL_BY_SOURCE_FOLDER[key]
        return model, display, note, False
    return "", "", "unmapped_source_folder", True


def load_authoritative_mapping(path: Path) -> dict[tuple[str, str], dict[str, str]]:
    rows = read_csv_rows(path)
    mapping: dict[tuple[str, str], dict[str, str]] = {}
    for row in rows:
        source_folder = (row.get("source_folder") or "").strip()
        source_file = (row.get("source_file") or "").replace("\\", "/").strip()
        if source_folder and source_file:
            mapping[(source_folder, source_file)] = row
    return mapping


def find_sample_ids(text: str) -> tuple[str, str]:
    text = text.replace("\\", "/")
    match = re.search(r"(?<!\d)(\d{2})-[^\n/]*?(?:__vs__|vs)(\d{2})-", text, flags=re.I)
    if match:
        return match.group(1), match.group(2)

    ids = re.findall(r"(?<!\d)(\d{2})-", text)
    if len(ids) >= 2:
        return ids[0], ids[1]
    if len(ids) == 1:
        return ids[0], ""
    return "", ""


def infer_replicate_id(path: Path, source_folder: str) -> str:
    combined = f"{path.parent.name}/{path.stem}"
    match = re.search(r"__rep0?(\d+)\b", combined, flags=re.I)
    if match:
        return f"rep{int(match.group(1)):02d}"

    if source_folder.strip().lower() == "chat gpt":
        match = re.search(r"-(\d+)$", path.stem)
        if match:
            return f"rep{int(match.group(1)):02d}"
    return ""


def infer_source_meta(path: Path, source_root: Path, authoritative: dict[tuple[str, str], dict[str, str]]) -> SourceMeta:
    rel = path.relative_to(source_root)
    source_folder = rel.parts[0] if rel.parts else ""
    source_file = Path(*rel.parts[1:]).as_posix() if len(rel.parts) > 1 else path.name

    model, display, note, uncertain = infer_model(source_folder)
    mapped = authoritative.get((source_folder, source_file))
    if mapped:
        model = (mapped.get("model") or model).strip()
        display = (mapped.get("model_display_name") or display).strip()
        note = (mapped.get("notes") or "authoritative_mapping").strip()
        uncertain = not bool(model)

    if source_folder.strip().lower() == "dis gpt old":
        model = "dis-gpt"
        display = "DIS GPT"
        note = "explicit_user_mapping_dis_gpt"
        uncertain = False

    if source_folder.strip().lower() == "dis gpt old" and re.fullmatch(r"\d+", path.stem):
        source_id = path.stem
        comparison_id = source_id
        replicate_id = ""
        run_id = source_id
    else:
        left_id, right_id = find_sample_ids(f"{source_file}/{path.parent.name}")
        source_id = right_id or left_id or ""
        comparison_id = f"{left_id}_vs_{right_id}" if left_id and right_id else source_id or path.stem
        replicate_id = infer_replicate_id(path, source_folder)
        run_id = f"{comparison_id}_{replicate_id}" if replicate_id else comparison_id

    if mapped:
        source_id = (mapped.get("source_id") or source_id).strip()
        run_id = (mapped.get("run_id") or run_id).strip()

    return SourceMeta(
        source_folder=source_folder,
        source_file=source_file,
        source_id=source_id,
        comparison_id=comparison_id,
        run_id=run_id,
        replicate_id=replicate_id,
        model=model,
        model_display_name=display,
        mapping_note=note,
        mapping_uncertain=uncertain,
    )


JSON_KEY_TO_FIELD = {
    field: field for field in DIMENSION_FIELDS
}
JSON_KEY_TO_FIELD.update(
    {
        "surface_texture": "surface_texture_change",
        "texture_change": "surface_texture_change",
        "area_size_change": "volume_change",
        "area_change": "volume_change",
        "size_change": "volume_change",
        "dissolution_speed": "dissolution_speed_time",
        "disintegration_speed": "dissolution_speed_time",
        "dissolution_speed_and_time": "dissolution_speed_time",
        "disintegration_dissolution_speed": "dissolution_speed_time",
        "dissolution_medium_characteristics": "dissolution_medium",
        "medium_characteristics": "dissolution_medium",
        "fragment_distribution": "fragment_distribution_density",
        "fragment_distribution_with_density": "fragment_distribution_density",
    }
)

DIMENSION_KEY_RE = re.compile(
    r'["\']?('
    + "|".join(re.escape(key) for key in sorted(JSON_KEY_TO_FIELD, key=len, reverse=True))
    + r')["\']?\s*"?\s*:\s*',
    flags=re.I,
)
STOP_KEY_RE = re.compile(
    r'["\']?(notes_for_later_summary|group_id|time_range_h|time_h|final_report|final judgment)["\']?\s*:',
    flags=re.I,
)

HEADING_TO_FIELD = {
    "color change": "color_change",
    "shape change": "shape_change",
    "surface texture": "surface_texture_change",
    "surface texture change": "surface_texture_change",
    "volume change": "volume_change",
    "area/size change": "volume_change",
    "dissolution speed and time": "dissolution_speed_time",
    "dissolution speed/time": "dissolution_speed_time",
    "dissolution speed": "dissolution_speed_time",
    "disintegration speed": "dissolution_speed_time",
    "physical state change": "physical_state_change",
    "dissolution medium": "dissolution_medium",
    "dissolution medium effect": "dissolution_medium",
    "fragment distribution": "fragment_distribution_density",
    "fragment distribution and density": "fragment_distribution_density",
    "fragment distribution with density": "fragment_distribution_density",
}

HEADING_RE = re.compile(
    r"(?im)^\s*(?:\d+\s*[\.)]\s*)?"
    r"(color change|shape change|surface texture change|surface texture|volume change|area/size change|"
    r"dissolution speed and time|dissolution speed/time|dissolution speed|disintegration speed|"
    r"physical state change|dissolution medium effect|dissolution medium|"
    r"fragment distribution and density|fragment distribution with density|fragment distribution)\s*$"
)


def clean_extracted_value(value: str) -> str:
    value = value.strip()
    value = re.sub(r'^\s*["\':]+', "", value)
    value = re.sub(r"\s+", " ", value)
    value = value.strip()
    while value.endswith((",", "}", "]")):
        value = value[:-1].strip()
    if value.endswith('"'):
        value = value[:-1].strip()
    if value.startswith('"'):
        value = value[1:].strip()
    return value.replace('\\"', '"').strip()


def extract_json_like_dimensions(text: str) -> dict[str, list[str]]:
    normalized = normalize_quotes(text)
    matches = list(DIMENSION_KEY_RE.finditer(normalized))
    found: dict[str, list[str]] = defaultdict(list)
    if not matches:
        return found

    for index, match in enumerate(matches):
        field = JSON_KEY_TO_FIELD.get(match.group(1).lower(), match.group(1).lower())
        start = match.end()
        end_candidates = []
        if index + 1 < len(matches):
            end_candidates.append(matches[index + 1].start())
        stop = STOP_KEY_RE.search(normalized, pos=start)
        if stop:
            end_candidates.append(stop.start())
        end = min(end_candidates) if end_candidates else len(normalized)
        value = clean_extracted_value(normalized[start:end])
        if value:
            found[field].append(value)
    return found


def extract_heading_dimensions(text: str) -> dict[str, list[str]]:
    normalized = normalize_quotes(text)
    matches = list(HEADING_RE.finditer(normalized))
    found: dict[str, list[str]] = defaultdict(list)
    if not matches:
        return found

    for index, match in enumerate(matches):
        heading = re.sub(r"\s+", " ", match.group(1).lower()).strip()
        field = HEADING_TO_FIELD.get(heading)
        if not field:
            continue
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(normalized)
        value = normalized[start:end]
        cutoff = re.search(
            r"(?im)^\s*(why i think|why these groups|final judgment|summary|overall conclusion|conclusion)\b",
            value,
        )
        if cutoff:
            value = value[: cutoff.start()]
        value = clean_extracted_value(value)
        if value:
            found[field].append(value)
    return found


def merge_dimension_hits(*hits: dict[str, list[str]]) -> dict[str, str]:
    merged: dict[str, str] = {}
    for field in DIMENSION_FIELDS:
        chunks: list[str] = []
        seen: set[str] = set()
        for hit in hits:
            for value in hit.get(field, []):
                normalized = re.sub(r"\s+", " ", value).strip()
                if normalized and normalized not in seen:
                    chunks.append(value)
                    seen.add(normalized)
        merged[field] = "\n\n---\n\n".join(chunks)
    return merged


def extract_dimensions(text: str) -> dict[str, str]:
    json_hits = extract_json_like_dimensions(text)
    heading_hits = extract_heading_dimensions(text)
    return merge_dimension_hits(json_hits, heading_hits)


def semantic_text(row: dict[str, object]) -> str:
    chunks = [str(row.get(field) or "").strip() for field in DIMENSION_FIELDS]
    text = "\n".join(chunk for chunk in chunks if chunk)
    return text or str(row.get("output_text") or "")


def clip_text(text: str, max_chars: int) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    if max_chars and len(text) > max_chars:
        return text[:max_chars]
    return text


def lexical_similarity(left: str, right: str) -> float:
    if not left.strip() and not right.strip():
        return 1.0
    if not left.strip() or not right.strip():
        return 0.0
    return SequenceMatcher(None, left, right).ratio()


def bertscore_batch(
    candidates: list[str],
    references: list[str],
    lang: str,
    model_type: str,
) -> list[float]:
    from bert_score import score

    _, _, f1 = score(
        candidates,
        references,
        lang=lang,
        model_type=model_type,
        verbose=False,
        rescale_with_baseline=False,
    )
    return [float(value) for value in f1.tolist()]


def load_reference_map(path: Path | None) -> dict[str, str]:
    if not path:
        return {}
    rows = read_csv_rows(path)
    mapping: dict[str, str] = {}
    for row in rows:
        text = (row.get("reference_text") or "").strip()
        if not text:
            continue
        for key_col in ("comparison_id", "source_id", "run_id", "reference_key"):
            key = (row.get(key_col) or "").strip()
            if key:
                mapping[key] = text
    return mapping


def apply_reference_bertscore(
    rows: list[dict[str, object]],
    reference_map: dict[str, str],
    args: argparse.Namespace,
) -> tuple[bool, str]:
    if not reference_map:
        for row in rows:
            row["bertscore_f1"] = ""
            row["reference_type"] = "reference_missing"
        return False, "No predefined reference file was provided."

    candidates: list[str] = []
    references: list[str] = []
    scored_indexes: list[int] = []
    for idx, row in enumerate(rows):
        reference = (
            reference_map.get(str(row.get("comparison_id") or ""))
            or reference_map.get(str(row.get("source_id") or ""))
            or reference_map.get(str(row.get("run_id") or ""))
        )
        if reference:
            candidates.append(clip_text(semantic_text(row), args.bert_max_chars))
            references.append(clip_text(reference, args.bert_max_chars))
            scored_indexes.append(idx)

    for row in rows:
        row["bertscore_f1"] = ""
        row["reference_type"] = "reference_missing"

    if not candidates:
        return False, "Reference file was provided, but no row matched comparison_id/source_id/run_id."

    try:
        scores = bertscore_batch(candidates, references, args.bert_lang, args.bert_model)
        for idx, score_value in zip(scored_indexes, scores, strict=True):
            rows[idx]["bertscore_f1"] = f"{score_value:.6f}"
            rows[idx]["reference_type"] = "predefined_reference_answer"
        return True, "BERTScore was calculated for rows with matched references."
    except Exception as exc:
        for idx in scored_indexes:
            rows[idx]["reference_type"] = "reference_score_failed"
        return False, f"BERTScore reference scoring failed: {exc}"


def build_pairwise_rows(
    rows: list[dict[str, object]],
    args: argparse.Namespace,
    allow_bertscore: bool,
) -> tuple[list[dict[str, object]], str]:
    groups: dict[tuple[str, str], list[dict[str, object]]] = defaultdict(list)
    for row in rows:
        model = str(row.get("model") or "")
        comparison_id = str(row.get("comparison_id") or "")
        if model and comparison_id:
            groups[(model, comparison_id)].append(row)

    pair_records: list[tuple[dict[str, object], dict[str, object]]] = []
    for group_rows in groups.values():
        unique_by_run: dict[str, dict[str, object]] = {}
        for row in group_rows:
            unique_by_run[str(row.get("run_id") or row.get("source_file") or "")] = row
        if len(unique_by_run) < 2:
            continue
        pair_records.extend(combinations(unique_by_run.values(), 2))

    if not pair_records:
        return [], "No repeated runs were available for pairwise similarity."

    candidates = [clip_text(semantic_text(left), args.bert_max_chars) for left, _ in pair_records]
    references = [clip_text(semantic_text(right), args.bert_max_chars) for _, right in pair_records]
    method = "lexical_fallback"
    scores: list[float]
    warning = ""

    if not allow_bertscore and args.pairwise_method in {"auto", "bertscore"}:
        warning = "BERTScore skipped because no predefined reference answer was provided; lexical similarity used."
        scores = [lexical_similarity(left, right) for left, right in zip(candidates, references, strict=True)]
        method = "lexical_no_reference"
    elif args.pairwise_method in {"auto", "bertscore"}:
        try:
            scores = bertscore_batch(candidates, references, args.bert_lang, args.bert_model)
            method = f"bertscore_f1:{args.bert_model}"
        except Exception as exc:
            if args.pairwise_method == "bertscore":
                raise
            warning = f"BERTScore pairwise scoring failed; lexical fallback used: {exc}"
            scores = [lexical_similarity(left, right) for left, right in zip(candidates, references, strict=True)]
    else:
        warning = "Lexical fallback requested explicitly."
        scores = [lexical_similarity(left, right) for left, right in zip(candidates, references, strict=True)]

    pairwise_rows: list[dict[str, object]] = []
    for (left, right), score_value in zip(pair_records, scores, strict=True):
        pairwise_rows.append(
            {
                "model": left.get("model", ""),
                "model_display_name": left.get("model_display_name", ""),
                "drug": "nifedipine",
                "comparison_id": left.get("comparison_id", ""),
                "run_id_1": left.get("run_id", ""),
                "run_id_2": right.get("run_id", ""),
                "source_file_1": left.get("source_file", ""),
                "source_file_2": right.get("source_file", ""),
                "pairwise_similarity": f"{float(score_value):.6f}",
                "similarity_method": method,
                "notes": warning,
            }
        )
    return pairwise_rows, warning or f"Pairwise similarity calculated with {method}."


def qc_flag(parse_status: str, missing_dimensions: list[str], mapping_uncertain: bool, reference_type: str) -> str:
    flags = []
    if parse_status != "read_ok":
        flags.append("parse_failed")
    if mapping_uncertain:
        flags.append("model_mapping_uncertain")
    if missing_dimensions:
        flags.append("missing_dimensions")
    if reference_type != "predefined_reference_answer":
        flags.append("reference_not_scored")
    return ";".join(flags) if flags else "ok"


def build_summary(
    source_root: Path,
    docx_files: list[Path],
    run_rows: list[dict[str, object]],
    qc_rows: list[dict[str, object]],
    pairwise_rows: list[dict[str, object]],
    reference_message: str,
    pairwise_message: str,
) -> str:
    top_folders = sorted({path.relative_to(source_root).parts[0] for path in docx_files if path.relative_to(source_root).parts})
    success_count = sum(1 for row in qc_rows if row.get("parse_status") == "read_ok")
    included_count = len(run_rows)
    model_counts = Counter(str(row.get("model_display_name") or row.get("model") or "") for row in run_rows)
    complete_8 = sum(1 for row in run_rows if int(row.get("completed_dimensions") or 0) == len(DIMENSION_FIELDS))
    expected = {"GPT-5-mini", "Qwen3.6-plus", "GLM-4.6V", "Kimi-K2.5"}
    detected = set(model_counts)
    expected_found = expected.issubset(detected)

    lines = [
        "# Fig. 3 Preprocessing QC Summary",
        "",
        f"- Source root: `{source_root}`",
        f"- Top-level folders detected: {len(top_folders)} ({', '.join(top_folders)})",
        f"- Word documents detected: {len(docx_files)}",
        f"- Documents successfully read: {success_count}",
        f"- Documents included for Fig. 3: {included_count}",
        f"- Pairwise similarity rows: {len(pairwise_rows)}",
        f"- Rows with all 8 dimensions extracted: {complete_8}/{included_count}",
        f"- All four expected plotting models found: {'yes' if expected_found else 'no'}",
        f"- Reference BERTScore status: {reference_message}",
        f"- Pairwise similarity status: {pairwise_message}",
        "",
        "## Run Count By Model",
        "",
        "| Model | Runs |",
        "| --- | ---: |",
    ]
    for model, count in sorted(model_counts.items()):
        lines.append(f"| {model or 'unmapped'} | {count} |")
    return "\n".join(lines) + "\n"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source_root", type=Path, help="Root folder containing model output folders.")
    parser.add_argument("--out-root", type=Path, default=ROOT, help="nature_comm_figures project root.")
    parser.add_argument("--reference-file", type=Path, default=None, help="Optional CSV with reference_text by comparison_id/source_id/run_id.")
    parser.add_argument(
        "--pairwise-method",
        choices=["auto", "bertscore", "lexical"],
        default="auto",
        help="Similarity method for repeated-run consistency.",
    )
    parser.add_argument("--bert-model", default="distilbert-base-uncased", help="Transformer backend for BERTScore.")
    parser.add_argument("--bert-lang", default="en", help="BERTScore language code.")
    parser.add_argument("--bert-max-chars", type=int, default=3500, help="Maximum chars per text passed to BERTScore.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    source_root = args.source_root.resolve()
    out_root = args.out_root.resolve()

    raw_runs_path = out_root / "data" / "raw" / "fig3_model_runs.csv"
    raw_pairwise_path = out_root / "data" / "raw" / "fig3_pairwise_similarity.csv"
    long_path = out_root / "data" / "processed" / "fig3_model_runs_long_by_dimension.csv"
    qc_path = out_root / "outputs" / "qc" / "fig3_qc_report.csv"
    mapping_path = out_root / "data" / "raw" / "fig3_file_mapping.csv"
    summary_path = out_root / "outputs" / "qc" / "fig3_preprocess_summary.md"

    authoritative = load_authoritative_mapping(mapping_path)
    docx_files = sorted(source_root.rglob("*.docx"))
    run_rows: list[dict[str, object]] = []
    qc_rows: list[dict[str, object]] = []
    mapping_rows: list[dict[str, object]] = []

    for path in docx_files:
        meta = infer_source_meta(path, source_root, authoritative)
        output_text, parse_status, parse_warning = read_docx(path)
        dimensions = extract_dimensions(output_text) if output_text else {field: "" for field in DIMENSION_FIELDS}
        completed = sum(1 for field in DIMENSION_FIELDS if dimensions.get(field, "").strip())
        completion_rate = completed / len(DIMENSION_FIELDS)
        missing = [field for field in DIMENSION_FIELDS if not dimensions.get(field, "").strip()]
        include = parse_status == "read_ok" and bool(meta.model)
        warning_parts = [part for part in [parse_warning, meta.mapping_note if meta.mapping_uncertain else ""] if part]

        base_row: dict[str, object] = {
            "model": meta.model,
            "model_display_name": meta.model_display_name,
            "drug": "nifedipine",
            "run_id": meta.run_id,
            "comparison_id": meta.comparison_id,
            "source_folder": meta.source_folder,
            "source_file": meta.source_file,
            "source_id": meta.source_id,
            "completed_dimensions": completed,
            "completion_rate": f"{completion_rate:.6f}",
            "bertscore_f1": "",
            "output_text": output_text,
            "reference_type": "",
            "notes": meta.mapping_note,
            **dimensions,
        }
        if include:
            run_rows.append(base_row)

        mapping_rows.append(
            {
                "source_folder": meta.source_folder,
                "source_file": meta.source_file,
                "model": meta.model,
                "model_display_name": meta.model_display_name,
                "run_id": meta.run_id,
                "source_id": meta.source_id,
                "include_for_fig3": int(include),
                "notes": meta.mapping_note,
            }
        )

        qc_rows.append(
            {
                "model": meta.model,
                "model_display_name": meta.model_display_name,
                "drug": "nifedipine",
                "run_id": meta.run_id,
                "source_folder": meta.source_folder,
                "source_file": meta.source_file,
                "source_id": meta.source_id,
                "parse_status": parse_status,
                "completed_dimensions": completed,
                "completion_rate": f"{completion_rate:.6f}",
                "missing_dimensions": ";".join(missing),
                "bertscore_f1": "",
                "output_text_length": len(output_text),
                "qc_flag": "",
                "warning_message": "; ".join(warning_parts),
            }
        )

    reference_map = load_reference_map(args.reference_file)
    reference_scored, reference_message = apply_reference_bertscore(run_rows, reference_map, args)
    bertscore_by_run = {str(row.get("run_id") or ""): row.get("bertscore_f1", "") for row in run_rows}
    reference_type_by_run = {str(row.get("run_id") or ""): row.get("reference_type", "") for row in run_rows}
    for qc_row in qc_rows:
        run_id = str(qc_row.get("run_id") or "")
        qc_row["bertscore_f1"] = bertscore_by_run.get(run_id, "")
        reference_type = reference_type_by_run.get(run_id, "reference_missing")
        missing_dimensions = [value for value in str(qc_row.get("missing_dimensions") or "").split(";") if value]
        qc_row["qc_flag"] = qc_flag(
            str(qc_row.get("parse_status") or ""),
            missing_dimensions,
            any(
                token in str(qc_row.get("warning_message") or "")
                for token in ("model_mapping_uncertain", "unmapped_source_folder")
            ),
            reference_type,
        )
        warnings = [str(qc_row.get("warning_message") or "").strip()]
        if reference_type != "predefined_reference_answer":
            warnings.append(reference_type)
        qc_row["warning_message"] = "; ".join(part for part in warnings if part)

    try:
        pairwise_rows, pairwise_message = build_pairwise_rows(
            run_rows,
            args,
            allow_bertscore=bool(reference_map),
        )
    except Exception:
        pairwise_rows = []
        pairwise_message = "Pairwise similarity failed:\n" + traceback.format_exc()

    long_rows: list[dict[str, object]] = []
    for row in run_rows:
        for field in DIMENSION_FIELDS:
            content = str(row.get(field) or "")
            long_rows.append(
                {
                    "model": row.get("model", ""),
                    "model_display_name": row.get("model_display_name", ""),
                    "drug": "nifedipine",
                    "run_id": row.get("run_id", ""),
                    "source_folder": row.get("source_folder", ""),
                    "source_file": row.get("source_file", ""),
                    "dimension": field,
                    "dimension_display_name": DIMENSION_DISPLAY.get(field, field),
                    "present": int(bool(content.strip())),
                    "content": content,
                    "content_length": len(content),
                }
            )

    write_csv(raw_runs_path, run_rows, RUN_COLUMNS)
    write_csv(raw_pairwise_path, pairwise_rows, PAIRWISE_COLUMNS)
    write_csv(long_path, long_rows, LONG_COLUMNS)
    write_csv(qc_path, qc_rows, QC_COLUMNS)
    write_csv(mapping_path, mapping_rows, MAPPING_COLUMNS)

    summary = build_summary(
        source_root,
        docx_files,
        run_rows,
        qc_rows,
        pairwise_rows,
        reference_message,
        pairwise_message,
    )
    ensure_parent(summary_path).write_text(summary, encoding="utf-8")

    print(summary)
    print("Outputs:")
    print(f"- {raw_runs_path}")
    print(f"- {raw_pairwise_path}")
    print(f"- {long_path}")
    print(f"- {qc_path}")
    print(f"- {mapping_path}")

    if not reference_scored:
        print("Reference BERTScore was not calculated; see QC for reference_missing rows.")


if __name__ == "__main__":
    main()
